# budget.py — PamojaData Budget Tracking Module
# Monitors programme budget utilisation alongside indicator performance
# Tracks expenditure, burn rate, variance and donor compliance

import sqlite3
import pandas as pd
import os
from datetime import datetime

DB_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'pamojadata.db')


def get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def initialise_budget_tables():
    """Creates budget tracking tables in the database."""
    conn = get_connection()
    cursor = conn.cursor()

    # Budget lines table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS budget_lines (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            programme_name TEXT NOT NULL,
            donor TEXT,
            budget_line TEXT NOT NULL,
            category TEXT,
            total_budget REAL DEFAULT 0,
            q1_budget REAL DEFAULT 0,
            q2_budget REAL DEFAULT 0,
            q3_budget REAL DEFAULT 0,
            q4_budget REAL DEFAULT 0,
            currency TEXT DEFAULT 'USD',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Expenditure table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS expenditure (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            budget_line_id INTEGER,
            amount REAL NOT NULL,
            description TEXT,
            reporting_period TEXT,
            expenditure_date TEXT,
            approved_by TEXT,
            receipt_reference TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (budget_line_id) REFERENCES budget_lines(id)
        )
    ''')

    conn.commit()
    conn.close()


def add_budget_line(programme_name, donor, budget_line, category,
                    total_budget, q1_budget, q2_budget,
                    q3_budget, q4_budget, currency='USD'):
    """Adds a new budget line."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO budget_lines (
            programme_name, donor, budget_line, category,
            total_budget, q1_budget, q2_budget,
            q3_budget, q4_budget, currency
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (programme_name, donor, budget_line, category,
          total_budget, q1_budget, q2_budget,
          q3_budget, q4_budget, currency))
    conn.commit()
    conn.close()


def add_expenditure(budget_line_id, amount, description,
                    reporting_period, expenditure_date,
                    approved_by, receipt_reference):
    """Records an expenditure against a budget line."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO expenditure (
            budget_line_id, amount, description,
            reporting_period, expenditure_date,
            approved_by, receipt_reference
        ) VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (budget_line_id, amount, description,
          reporting_period, expenditure_date,
          approved_by, receipt_reference))
    conn.commit()
    conn.close()


def get_budget_summary(programme_name=None):
    """
    Returns budget vs expenditure summary per budget line.
    Calculates utilisation rate and variance.
    """
    conn = get_connection()
    cursor = conn.cursor()

    if programme_name:
        cursor.execute('''
            SELECT
                bl.id,
                bl.programme_name,
                bl.donor,
                bl.budget_line,
                bl.category,
                bl.total_budget,
                bl.currency,
                COALESCE(SUM(e.amount), 0) as total_spent
            FROM budget_lines bl
            LEFT JOIN expenditure e ON bl.id = e.budget_line_id
            WHERE bl.programme_name = ?
            GROUP BY bl.id
        ''', (programme_name,))
    else:
        cursor.execute('''
            SELECT
                bl.id,
                bl.programme_name,
                bl.donor,
                bl.budget_line,
                bl.category,
                bl.total_budget,
                bl.currency,
                COALESCE(SUM(e.amount), 0) as total_spent
            FROM budget_lines bl
            LEFT JOIN expenditure e ON bl.id = e.budget_line_id
            GROUP BY bl.id
        ''')

    rows = [dict(row) for row in cursor.fetchall()]
    conn.close()

    if not rows:
        return pd.DataFrame()

    df = pd.DataFrame(rows)
    df['Variance'] = df['total_budget'] - df['total_spent']
    df['Utilisation %'] = df.apply(
        lambda row: round(row['total_spent'] / row['total_budget'] * 100, 1)
        if row['total_budget'] > 0 else 0, axis=1
    )
    df['Status'] = df['Utilisation %'].apply(
        lambda x: '🔴 Over Budget' if x > 100
        else '🟡 High Burn' if x > 80
        else '✅ On Track' if x >= 40
        else '⚠️ Under Spent'
    )
    return df


def get_expenditure_by_period(programme_name=None):
    """Returns expenditure grouped by reporting period."""
    conn = get_connection()
    cursor = conn.cursor()

    query = '''
        SELECT
            e.reporting_period,
            SUM(e.amount) as total_spent,
            COUNT(e.id) as transactions
        FROM expenditure e
        JOIN budget_lines bl ON e.budget_line_id = bl.id
    '''
    params = []
    if programme_name:
        query += ' WHERE bl.programme_name = ?'
        params.append(programme_name)
    query += ' GROUP BY e.reporting_period ORDER BY e.reporting_period'

    cursor.execute(query, tuple(params))
    rows = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return pd.DataFrame(rows) if rows else pd.DataFrame()


def get_category_summary(programme_name=None):
    """Returns budget utilisation by category."""
    df = get_budget_summary(programme_name)
    if df.empty:
        return pd.DataFrame()

    return df.groupby('category').agg(
        Total_Budget=('total_budget', 'sum'),
        Total_Spent=('total_spent', 'sum')
    ).reset_index().assign(
        Utilisation=lambda x: (
            x['Total_Spent'] / x['Total_Budget'] * 100
        ).round(1)
    )


def get_overall_budget_kpis(programme_name=None):
    """Returns headline budget KPIs."""
    df = get_budget_summary(programme_name)
    if df.empty:
        return {
            'total_budget': 0,
            'total_spent': 0,
            'total_remaining': 0,
            'overall_utilisation': 0,
            'over_budget_lines': 0,
            'under_spent_lines': 0
        }

    total_budget = df['total_budget'].sum()
    total_spent = df['total_spent'].sum()

    return {
        'total_budget': total_budget,
        'total_spent': total_spent,
        'total_remaining': total_budget - total_spent,
        'overall_utilisation': round(
            total_spent / total_budget * 100, 1
        ) if total_budget > 0 else 0,
        'over_budget_lines': len(df[df['Utilisation %'] > 100]),
        'under_spent_lines': len(df[df['Utilisation %'] < 40])
    }


def get_all_programmes():
    """Returns list of all programmes with budget lines."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        'SELECT DISTINCT programme_name FROM budget_lines ORDER BY programme_name'
    )
    programmes = [row[0] for row in cursor.fetchall()]
    conn.close()
    return programmes


def get_budget_lines(programme_name=None):
    """Returns all budget lines."""
    conn = get_connection()
    cursor = conn.cursor()
    if programme_name:
        cursor.execute(
            'SELECT * FROM budget_lines WHERE programme_name = ?',
            (programme_name,)
        )
    else:
        cursor.execute('SELECT * FROM budget_lines')
    rows = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return rows


def export_budget_to_word(programme_name=None):
    """Exports budget report as Word document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    df = get_budget_summary(programme_name)
    kpis = get_overall_budget_kpis(programme_name)
    category_sum = get_category_summary(programme_name)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading('', level=0)
    run = title.add_run(
        f'Budget Utilisation Report — {programme_name or "All Programmes"}'
    )
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta.add_run(f'Generated: {datetime.now().strftime("%d %B %Y")}')
    doc.add_paragraph('─' * 80)

    # KPIs
    doc.add_heading('Budget Summary', level=1)
    kpi_para = doc.add_paragraph()
    kpi_para.add_run('Total Budget: ').bold = True
    kpi_para.add_run(f"${kpis['total_budget']:,.0f}   |   ")
    kpi_para.add_run('Total Spent: ').bold = True
    kpi_para.add_run(f"${kpis['total_spent']:,.0f}   |   ")
    kpi_para.add_run('Remaining: ').bold = True
    kpi_para.add_run(f"${kpis['total_remaining']:,.0f}   |   ")
    kpi_para.add_run('Utilisation: ').bold = True
    kpi_para.add_run(f"{kpis['overall_utilisation']}%")

    # Budget lines table
    if not df.empty:
        doc.add_heading('Budget Line Details', level=1)
        cols = ['budget_line', 'category', 'total_budget',
                'total_spent', 'Variance', 'Utilisation %', 'Status']
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'
        headers = ['Budget Line', 'Category', 'Total Budget',
                   'Spent', 'Variance', 'Utilisation %', 'Status']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                val = row[col]
                if col in ['total_budget', 'total_spent', 'Variance']:
                    cells[i].text = f"${val:,.0f}"
                else:
                    cells[i].text = str(val)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# Initialise on import
initialise_budget_tables()