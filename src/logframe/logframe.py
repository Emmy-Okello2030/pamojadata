# logframe.py — PamojaData Logframe Builder
# Creates and manages logical frameworks for humanitarian programmes
# Links Goals → Outcomes → Outputs → Activities with indicators and targets

import sqlite3
import pandas as pd
import os

DB_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'pamojadata.db')


def get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def initialise_logframe_tables():
    """
    Creates logframe-related tables in the database.
    """
    conn = get_connection()
    cursor = conn.cursor()

    # Logframe table — one per programme
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS logframes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            programme_id INTEGER,
            programme_name TEXT NOT NULL,
            org_name TEXT,
            donor TEXT,
            start_date TEXT,
            end_date TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Logframe entries — Goals, Outcomes, Outputs, Activities
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS logframe_entries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            logframe_id INTEGER,
            level TEXT NOT NULL,
            description TEXT NOT NULL,
            indicator TEXT,
            means_of_verification TEXT,
            assumptions TEXT,
            target REAL,
            baseline REAL,
            responsible_party TEXT,
            timeline TEXT,
            parent_id INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (logframe_id) REFERENCES logframes(id),
            FOREIGN KEY (parent_id) REFERENCES logframe_entries(id)
        )
    ''')

    conn.commit()
    conn.close()


def create_logframe(programme_name, org_name, donor, start_date, end_date):
    """Creates a new logframe record."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO logframes (programme_name, org_name, donor, start_date, end_date)
        VALUES (?, ?, ?, ?, ?)
    ''', (programme_name, org_name, donor, start_date, end_date))
    logframe_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return logframe_id


def add_entry(logframe_id: int, level: str, description: str, indicator: str = "",
              means_of_verification: str = "", assumptions: str = "",
              target: float = 0.0, baseline: float = 0.0, responsible_party: str = "",
              timeline: str = "", parent_id=None):
    """
    Adds a Goal, Outcome, Output or Activity entry to a logframe.
    parent_id links outputs to outcomes, activities to outputs.
    """
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO logframe_entries
        (logframe_id, level, description, indicator,
         means_of_verification, assumptions, target,
         baseline, responsible_party, timeline, parent_id)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (logframe_id, level, description, indicator,
          means_of_verification, assumptions, target,
          baseline, responsible_party, timeline, parent_id))
    entry_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return entry_id


def get_logframe(logframe_id):
    """Retrieves a full logframe with all entries."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute('SELECT * FROM logframes WHERE id = ?', (logframe_id,))
    logframe = dict(cursor.fetchone())

    cursor.execute('''
        SELECT * FROM logframe_entries
        WHERE logframe_id = ?
        ORDER BY
            CASE level
                WHEN 'Goal' THEN 1
                WHEN 'Outcome' THEN 2
                WHEN 'Output' THEN 3
                WHEN 'Activity' THEN 4
            END, id
    ''', (logframe_id,))
    entries = [dict(row) for row in cursor.fetchall()]
    conn.close()

    return logframe, entries


def get_all_logframes():
    """Retrieves all logframes for the sidebar list."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM logframes ORDER BY created_at DESC')
    rows = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return rows


def delete_entry(entry_id):
    """Deletes a single logframe entry."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM logframe_entries WHERE id = ?', (entry_id,))
    conn.commit()
    conn.close()


def delete_logframe(logframe_id):
    """Deletes a logframe and all its entries."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        'DELETE FROM logframe_entries WHERE logframe_id = ?',
        (logframe_id,)
    )
    cursor.execute('DELETE FROM logframes WHERE id = ?', (logframe_id,))
    conn.commit()
    conn.close()


def logframe_to_dataframe(logframe_id):
    """
    Converts a logframe to a pandas DataFrame.
    Used for export to Word and feeding into analysis module.
    """
    _, entries = get_logframe(logframe_id)
    if not entries:
        return pd.DataFrame()

    df = pd.DataFrame(entries)[[
        'level', 'description', 'indicator',
        'baseline', 'target', 'means_of_verification',
        'assumptions', 'responsible_party', 'timeline'
    ]]

    df.columns = [
        'Level', 'Description', 'Indicator',
        'Baseline', 'Target', 'Means of Verification',
        'Assumptions', 'Responsible Party', 'Timeline'
    ]

    # Order by level
    level_order = {'Goal': 1, 'Outcome': 2, 'Output': 3, 'Activity': 4}
    df['_order'] = df['Level'].map(level_order)
    df = df.sort_values('_order').drop('_order', axis=1)

    return df


def export_logframe_to_word(logframe_id):
    """
    Exports a logframe as a formatted Word document.
    Returns BytesIO buffer.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    from docx.oxml import parse_xml  # type: ignore[import]
    from datetime import datetime
    import io

    logframe, entries = get_logframe(logframe_id)
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading('', level=0)
    run = title.add_run(logframe['programme_name'])
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run.font.size = Pt(22)

    sub = doc.add_paragraph()
    sub.add_run('Logical Framework (Logframe)   |   ').bold = True
    sub.add_run(f"{logframe['org_name']}   |   ")
    sub.add_run(f"Donor: {logframe['donor']}")

    period = doc.add_paragraph()
    period.add_run('Programme Period: ').bold = True
    period.add_run(
        f"{logframe['start_date']} to {logframe['end_date']}"
        f"   |   Generated: {datetime.now().strftime('%d %B %Y')}"
    )

    doc.add_paragraph('─' * 80)
    doc.add_paragraph()

    # Level colors
    level_colors = {
        'Goal': '1A3C5E',
        'Outcome': '1A8A7A',
        'Output': '2E6DA4',
        'Activity': '5D6D7E'
    }

    # Group entries by level
    levels = ['Goal', 'Outcome', 'Output', 'Activity']
    grouped = {level: [] for level in levels}
    for entry in entries:
        if entry['level'] in grouped:
            grouped[entry['level']].append(entry)

    for level in levels:
        level_entries = grouped[level]
        if not level_entries:
            continue

        # Level heading
        heading = doc.add_heading('', level=1)
        h_run = heading.add_run(
            f"{'🎯' if level == 'Goal' else '📌' if level == 'Outcome' else '📦' if level == 'Output' else '⚙️'} {level}s"
        )
        h_run.font.color.rgb = RGBColor(
            *bytes.fromhex(level_colors[level])
        )

        # Table for this level
        headers = [
            'Description', 'Indicator', 'Baseline',
            'Target', 'Means of Verification',
            'Assumptions', 'Responsible', 'Timeline'
        ]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(
                0xFF, 0xFF, 0xFF
            )
            try:
                shading = parse_xml(
                    f'<w:shd '
                    f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                    f'w:fill="{level_colors[level]}" w:val="clear"/>'
                )
                hdr[i]._tc.get_or_add_tcPr().append(shading)
            except Exception:
                pass

        # Data rows
        for entry in level_entries:
            row = table.add_row().cells
            row[0].text = entry['description'] or ''
            row[1].text = entry['indicator'] or ''
            row[2].text = str(entry['baseline'] or 0)
            row[3].text = str(entry['target'] or 0)
            row[4].text = entry['means_of_verification'] or ''
            row[5].text = entry['assumptions'] or ''
            row[6].text = entry['responsible_party'] or ''
            row[7].text = entry['timeline'] or ''

            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9) \
                    if cell.paragraphs[0].runs else None

        doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph(
        f'PamojaData | {logframe["org_name"]} | '
        f'Generated {datetime.now().strftime("%d %B %Y")} | Confidential'
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# Initialise tables on import
initialise_logframe_tables()