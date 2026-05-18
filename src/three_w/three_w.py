# three_w.py — PamojaData 3W Tracking Module
# Who does What Where — standard humanitarian coordination tool
# Maps operational presence across organisations, sectors and locations
# Reference: OCHA 3W/4W Operational Presence Methodology

import sqlite3
import pandas as pd
import os
from datetime import datetime

DB_PATH = os.path.join(os.path.dirname(__file__), '..', '..', 'data', 'pamojadata.db')


def get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def initialise_three_w_tables():
    """Creates 3W tracking tables in the database."""
    conn = get_connection()
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS three_w (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            organisation TEXT NOT NULL,
            organisation_type TEXT,
            sector TEXT NOT NULL,
            subsector TEXT,
            activity TEXT NOT NULL,
            admin1 TEXT,
            admin2 TEXT,
            admin3 TEXT,
            location_name TEXT,
            latitude REAL,
            longitude REAL,
            beneficiaries_targeted INTEGER DEFAULT 0,
            beneficiaries_reached INTEGER DEFAULT 0,
            start_date TEXT,
            end_date TEXT,
            status TEXT DEFAULT 'Active',
            funding_source TEXT,
            contact_name TEXT,
            contact_email TEXT,
            reporting_period TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    conn.commit()
    conn.close()


def add_three_w_entry(organisation, organisation_type, sector,
                       subsector, activity, admin1, admin2,
                       admin3, location_name, latitude, longitude,
                       beneficiaries_targeted, beneficiaries_reached,
                       start_date, end_date, status, funding_source,
                       contact_name, contact_email, reporting_period):
    """Adds a new 3W entry."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO three_w (
            organisation, organisation_type, sector, subsector,
            activity, admin1, admin2, admin3, location_name,
            latitude, longitude, beneficiaries_targeted,
            beneficiaries_reached, start_date, end_date, status,
            funding_source, contact_name, contact_email, reporting_period
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        organisation, organisation_type, sector, subsector,
        activity, admin1, admin2, admin3, location_name,
        latitude, longitude, beneficiaries_targeted,
        beneficiaries_reached, start_date, end_date, status,
        funding_source, contact_name, contact_email, reporting_period
    ))
    conn.commit()
    conn.close()


def get_all_three_w(reporting_period=None):
    """Retrieves all 3W entries, optionally filtered by period."""
    conn = get_connection()
    cursor = conn.cursor()

    if reporting_period:
        cursor.execute(
            'SELECT * FROM three_w WHERE reporting_period = ? ORDER BY organisation',
            (reporting_period,)
        )
    else:
        cursor.execute('SELECT * FROM three_w ORDER BY organisation')

    rows = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return rows


def get_three_w_dataframe(reporting_period=None):
    """Returns 3W data as a pandas DataFrame."""
    rows = get_all_three_w(reporting_period)
    if not rows:
        return pd.DataFrame()
    return pd.DataFrame(rows)


def get_sector_presence(reporting_period=None):
    """Returns organisation count per sector."""
    df = get_three_w_dataframe(reporting_period)
    if df.empty:
        return pd.DataFrame()
    return df.groupby('sector').agg(
        Organisations=('organisation', 'nunique'),
        Activities=('activity', 'count'),
        Beneficiaries_Targeted=('beneficiaries_targeted', 'sum'),
        Beneficiaries_Reached=('beneficiaries_reached', 'sum')
    ).reset_index()


def get_location_presence(reporting_period=None):
    """Returns organisation count per location."""
    df = get_three_w_dataframe(reporting_period)
    if df.empty:
        return pd.DataFrame()
    return df.groupby('admin1').agg(
        Organisations=('organisation', 'nunique'),
        Sectors=('sector', 'nunique'),
        Activities=('activity', 'count'),
        Beneficiaries_Reached=('beneficiaries_reached', 'sum')
    ).reset_index()


def get_organisation_summary(reporting_period=None):
    """Returns activity summary per organisation."""
    df = get_three_w_dataframe(reporting_period)
    if df.empty:
        return pd.DataFrame()
    return df.groupby(['organisation', 'organisation_type']).agg(
        Sectors=('sector', 'nunique'),
        Locations=('admin1', 'nunique'),
        Activities=('activity', 'count'),
        Beneficiaries_Reached=('beneficiaries_reached', 'sum')
    ).reset_index()


def get_coverage_gaps(reporting_period=None):
    """
    Identifies sectors and locations with low coverage.
    Flags areas where fewer than 2 organisations are present.
    """
    loc_presence = get_location_presence(reporting_period)
    if loc_presence.empty:
        return pd.DataFrame()

    gaps = loc_presence[loc_presence['Organisations'] < 2].copy()
    gaps['Gap Type'] = gaps.apply(
        lambda row: '🔴 No Coverage' if row['Organisations'] == 0
        else '🟡 Single Organisation', axis=1
    )
    return gaps


def bulk_import_three_w(df, mapping):
    """
    Bulk imports 3W data from a DataFrame.
    mapping: dict mapping standard fields to df column names
    """
    conn = get_connection()
    cursor = conn.cursor()
    count = 0

    for _, row in df.iterrows():
        try:
            cursor.execute('''
                INSERT INTO three_w (
                    organisation, sector, activity,
                    admin1, beneficiaries_targeted,
                    beneficiaries_reached, reporting_period
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                row.get(mapping.get('organisation', ''), ''),
                row.get(mapping.get('sector', ''), ''),
                row.get(mapping.get('activity', ''), ''),
                row.get(mapping.get('admin1', ''), ''),
                row.get(mapping.get('beneficiaries_targeted', ''), 0),
                row.get(mapping.get('beneficiaries_reached', ''), 0),
                row.get(mapping.get('reporting_period', ''), '')
            ))
            count += 1
        except Exception:
            continue

    conn.commit()
    conn.close()
    return count


def delete_three_w_entry(entry_id):
    """Deletes a single 3W entry."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute('DELETE FROM three_w WHERE id = ?', (entry_id,))
    conn.commit()
    conn.close()


def get_reporting_periods():
    """Returns list of all reporting periods in the database."""
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        'SELECT DISTINCT reporting_period FROM three_w ORDER BY reporting_period DESC'
    )
    periods = [row[0] for row in cursor.fetchall() if row[0]]
    conn.close()
    return periods


def export_three_w_to_word(reporting_period=None):
    """
    Exports 3W data as a formatted Word document.
    Returns BytesIO buffer.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    df = get_three_w_dataframe(reporting_period)
    sector_presence = get_sector_presence(reporting_period)
    location_presence = get_location_presence(reporting_period)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading('', level=0)
    run = title.add_run('3W Operational Presence Matrix')
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run.font.size = Pt(22)

    meta = doc.add_paragraph()
    meta.add_run('Who does What Where   |   ').bold = True
    meta.add_run(f'Reporting Period: {reporting_period or "All Periods"}   |   ')
    meta.add_run(f'Generated: {datetime.now().strftime("%d %B %Y")}')
    doc.add_paragraph('─' * 80)

    # Summary stats
    if not df.empty:
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run(f'Total Organisations: ').bold = True
        summary_para.add_run(f"{df['organisation'].nunique()}   |   ")
        summary_para.add_run(f'Sectors Covered: ').bold = True
        summary_para.add_run(f"{df['sector'].nunique()}   |   ")
        summary_para.add_run(f'Locations: ').bold = True
        summary_para.add_run(f"{df['admin1'].nunique()}   |   ")
        summary_para.add_run(f'Total Beneficiaries Reached: ').bold = True
        summary_para.add_run(f"{df['beneficiaries_reached'].sum():,}")

    # Sector presence table
    if not sector_presence.empty:
        doc.add_heading('Sector Presence', level=1)
        table = doc.add_table(rows=1, cols=len(sector_presence.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(sector_presence.columns):
            table.rows[0].cells[i].text = col
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for _, row in sector_presence.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    # Location presence table
    if not location_presence.empty:
        doc.add_page_break()
        doc.add_heading('Location Presence', level=1)
        table = doc.add_table(rows=1, cols=len(location_presence.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(location_presence.columns):
            table.rows[0].cells[i].text = col
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for _, row in location_presence.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    # Full 3W matrix
    if not df.empty:
        doc.add_page_break()
        doc.add_heading('Full 3W Matrix', level=1)
        cols = ['organisation', 'sector', 'activity',
                'admin1', 'beneficiaries_reached', 'status']
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'
        headers = ['Organisation', 'Sector', 'Activity',
                   'Location', 'Beneficiaries Reached', 'Status']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                cells[i].text = str(row.get(col, ''))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# Initialise on import
initialise_three_w_tables()