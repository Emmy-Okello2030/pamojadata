# report_export.py — PamojaData Report Export Engine
# Assembles final Word document from narrative, tables and charts

from docx import Document
from docx.shared import Inches, Pt, RGBColor  # type: ignore[import]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore[import]
from datetime import datetime
import io


def set_cell_color(cell, hex_color):
    """Sets background color of a table cell."""
    from docx.oxml import parse_xml  # type: ignore[import]
    shading = parse_xml(
        f'<w:shd '
        f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def create_word_report(analyzed_df, sector_summary, narrative,
                        mapping, org_name="My NGO", report_period="",
                        chart_buf=None, trend_df=None,
                        disagg_df=None, logo=None):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    if logo:
        try:
            doc.add_picture(logo, width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    title = doc.add_heading('', level=0)
    run = title.add_run(org_name)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run.font.size = Pt(24)

    sub = doc.add_heading('', level=2)
    sub_run = sub.add_run('Donor Progress Report')
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.add_run('Reporting Period: ').bold = True
    meta.add_run(f'{report_period if report_period else "Current Period"}   |   ')
    meta.add_run('Generated: ').bold = True
    meta.add_run(datetime.now().strftime('%d %B %Y'))
    meta.paragraph_format.space_after = Pt(12)

    doc.add_paragraph('─' * 80)

    doc.add_heading('Programme Narrative', level=1)
    for line in narrative.split('\n'):
        if line.strip():
            if line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            else:
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
    doc.add_heading('Indicator Performance Table', level=1)

    headers = ['Indicator', 'Sector', 'Target',
               'Achieved', 'Variance', '% Achievement', 'Status']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(table.rows[0].cells):
        h.text = headers[i]
        h.paragraphs[0].runs[0].bold = True
        h.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        h.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            set_cell_color(h, '1A3C5E')
        except Exception:
            pass

    for _, row in analyzed_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row[mapping['indicator_name']])
        cells[1].text = str(row[mapping['sector']])
        cells[2].text = str(int(row[mapping['target']]))
        cells[3].text = str(int(row[mapping['achieved']]))
        cells[4].text = f"{row['Variance']:+.0f}"
        cells[5].text = f"{row['% Achievement']:.1f}%"
        cells[6].text = row['Status']

        status = row['Status']
        bg = ('D5F5E3' if 'Met' in status
              else 'FEF9E7' if 'On Track' in status
              else 'FADBD8')
        try:
            for cell in cells:
                set_cell_color(cell, bg)
        except Exception:
            pass

    if chart_buf:
        doc.add_page_break()
        doc.add_heading('Visual Progress Summary', level=1)
        try:
            doc.add_picture(chart_buf, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    if disagg_df is not None:
        doc.add_page_break()
        doc.add_heading('Geographic Breakdown', level=1)
        dg = doc.add_table(rows=1, cols=len(disagg_df.columns))
        dg.style = 'Table Grid'
        for i, col in enumerate(disagg_df.columns):
            cell = dg.rows[0].cells[i]
            cell.text = str(col)
            cell.paragraphs[0].runs[0].bold = True
        for _, row in disagg_df.iterrows():
            cells = dg.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    doc.add_paragraph()
    footer = doc.add_paragraph(
        f'PamojaData | {org_name} | '
        f'Generated {datetime.now().strftime("%d %B %Y")} | Confidential'
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf